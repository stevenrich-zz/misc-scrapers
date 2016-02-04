import csv
import re
import string
import sys
import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
sections = document.sections
for section in sections:
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

document.add_paragraph(file_date)
document.add_paragraph("")
document.add_paragraph(dept)
document.add_paragraph(address1)
document.add_paragraph(address2)
if address3 != "":
    document.add_paragraph(address3)
if address4 != "":
    document.add_paragraph(address4)
if address5 != "":
    document.add_paragraph(address5)
document.add_paragraph("")
document.add_paragraph("To whom it may concern:")
document.add_paragraph("")
if statute2 == "":
	  if statute3 != "":
		    document.add_paragraph("This is a request under " + str(statute1) + u'\xa7' + u'\xa7' + str(statute3) +" I hereby request disclosure of the following information related to each officer who fired shots in the " + str(shooting_date) + " fatal shooting of " + str(deceased) + ":")
if statute3 == "":
		document.add_paragraph("This is a request under " + str(statute1) + " I hereby request disclosure of the following information related to each officer who fired shots in the " + str(shooting_date) + " fatal shooting of " + str(deceased) + ":")
if statute2 != "":
    document.add_paragraph("This is a request under " + str(statute1) + u'\xa7' + str(statute2) +" I hereby request disclosure of the following information related to each officer who fired shots in the " + str(shooting_date) + " fatal shooting of " + str(deceased) + ":")

document.add_paragraph("If you regard any of the requested records as subject to exemption from required disclosure under the law, I hereby request that you exercise your discretion and disclose them nonetheless.")
document.add_paragraph("")
document.add_paragraph("If the reason you are unable to disclose information is that there is an ongoing investigation, please notify me and wait to fulfill the request until such time as the investigation closes.")
document.add_paragraph("")
document.add_paragraph("If you decide to withhold any requested records, please do not deny the entire request and release information available under the law. If you cannot fulfill certain parts of this request or certain portions must be redacted please provide any reasonable segregable portion of a requested record after removing or redacting those portions claimed to be exempt, explain in writing the justification for redacting the record, indicate the extent of any redactions on the portion of the record which is made available or published and where technically feasible indicate on the redacted portion itself the specific exemption(s) claimed.")
document.add_paragraph("")
document.add_paragraph("I make this request on behalf of The Washington Post, a newspaper of general circulation in the Washington D.C. metropolitan area. The records disclosed pursuant to this request will be used in the preparation of news articles for dissemination to the public. Accordingly, I request that you waive all fees in the public interest because the furnishing of the information sought by this request will primarily benefit the public and is likely to contribute significantly to public understanding of the operations or activities of the government. If, however, you decline to waive all fees, I am prepared to pay, but request that you notify me if you plan to charge.")
document.add_paragraph("")
document.add_paragraph("If you have any questions about this request, please do not hesitate to contact me at policefoia@washpost.com or at the above telephone number. I look forward to your response.")
document.add_paragraph("")
document.add_paragraph("Sincerely,")
